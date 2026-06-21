import os
import json
from typing import List, Dict, Optional, Any
from pydantic import BaseModel, Field
from datetime import datetime
import uuid


class WorkflowNode(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4())[:8])
    type: str  # start, llm, tool, condition, transform, output
    label: str = ""
    config: Dict[str, Any] = {}
    position: Dict[str, float] = {"x": 0, "y": 0}


class WorkflowEdge(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4())[:8])
    source: str
    target: str
    sourceHandle: Optional[str] = None


class Workflow(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str = "未命名工作流"
    description: str = ""
    nodes: List[WorkflowNode] = []
    edges: List[WorkflowEdge] = []
    status: str = "draft"
    created_at: str = Field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = Field(default_factory=lambda: datetime.now().isoformat())


class WorkflowRun(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    workflow_id: str
    status: str = "running"
    inputs: Dict[str, Any] = {}
    outputs: Dict[str, Any] = {}
    node_results: Dict[str, Any] = {}
    started_at: str = Field(default_factory=lambda: datetime.now().isoformat())
    finished_at: Optional[str] = None
    error: Optional[str] = None


_workflows_col = None
_runs_col = None


def _get_cols():
    global _workflows_col, _runs_col
    if _workflows_col is None:
        from pymongo import MongoClient
        MONGO_URL = os.environ.get("MONGODB_URL", "mongodb://ai_mongo:ai_mongo_2024@localhost:27017")
        client = MongoClient(MONGO_URL, serverSelectionTimeoutMS=3000)
        db = client["ai_framework"]
        _workflows_col = db["workflows"]
        _runs_col = db["workflow_runs"]
    return _workflows_col, _runs_col


def list_workflows() -> List[Dict]:
    col, _ = _get_cols()
    results = []
    for doc in col.find().sort("updated_at", -1):
        doc["id"] = str(doc.pop("_id"))
        results.append(doc)
    return results


def get_workflow(workflow_id: str) -> Optional[Dict]:
    col, _ = _get_cols()
    doc = col.find_one({"_id": workflow_id})
    if doc:
        doc["id"] = str(doc.pop("_id"))
    return doc


def create_workflow(data: dict) -> Dict:
    col, _ = _get_cols()
    wf = Workflow(**data)
    doc = wf.model_dump()
    doc["_id"] = doc.pop("id")
    col.insert_one(doc)
    doc["id"] = doc["_id"]
    return doc


def update_workflow(workflow_id: str, data: dict) -> bool:
    col, _ = _get_cols()
    data["updated_at"] = datetime.now().isoformat()
    result = col.update_one({"_id": workflow_id}, {"$set": data})
    return result.modified_count > 0


def delete_workflow(workflow_id: str) -> bool:
    col, _ = _get_cols()
    result = col.delete_one({"_id": workflow_id})
    return result.deleted_count > 0


def save_run(run: WorkflowRun):
    _, runs_col = _get_cols()
    doc = {
        "_id": run.id,
        "workflow_id": run.workflow_id,
        "status": run.status,
        "inputs": run.inputs,
        "outputs": {},
        "node_results": {},
        "started_at": run.started_at,
        "finished_at": run.finished_at,
        "error": run.error
    }
    runs_col.insert_one(doc)


def update_run(run: WorkflowRun):
    _, runs_col = _get_cols()
    data = run.model_dump()
    data.pop("id", None)
    runs_col.update_one({"_id": run.id}, {"$set": {"status": data["status"], "outputs": data["outputs"], "node_results": data["node_results"], "finished_at": data.get("finished_at"), "error": data.get("error")}})


def get_runs(workflow_id: str, limit: int = 20) -> List[Dict]:
    _, runs_col = _get_cols()
    results = []
    for doc in runs_col.find({"workflow_id": workflow_id}).sort("started_at", -1).limit(limit):
        doc["id"] = str(doc.pop("_id"))
        results.append(doc)
    return results


def get_run(run_id: str) -> Optional[Dict]:
    _, runs_col = _get_cols()
    doc = runs_col.find_one({"_id": run_id})
    if doc:
        doc["id"] = str(doc.pop("_id"))
    return doc


def _topological_sort(nodes: List[dict], edges: List[dict]) -> List[List[str]]:
    in_degree = {n["id"]: 0 for n in nodes}
    children = {n["id"]: [] for n in nodes}
    for e in edges:
        in_degree[e["target"]] = in_degree.get(e["target"], 0) + 1
        children[e["source"]].append(e["target"])

    levels = []
    queue = [nid for nid, deg in in_degree.items() if deg == 0]
    while queue:
        levels.append(queue[:])
        next_queue = []
        for nid in queue:
            for child in children.get(nid, []):
                in_degree[child] -= 1
                if in_degree[child] == 0:
                    next_queue.append(child)
        queue = next_queue
    return levels


def _resolve_var(name: str, context: dict) -> str:
    """从 context 中解析变量，支持 {{key}} 和 {{node_id.field}} 两种格式"""
    def _extract(val):
        if isinstance(val, dict):
            if "response" in val:
                return str(val["response"])
            if "result" in val:
                return str(val["result"])
            if "stdout" in val:
                return str(val["stdout"])
            if "output" in val:
                inner = val["output"]
                if isinstance(inner, dict) and "response" in inner:
                    return str(inner["response"])
                if isinstance(inner, dict) and "result" in inner:
                    return str(inner["result"])
                if isinstance(inner, dict) and "stdout" in inner:
                    return str(inner["stdout"])
                return json.dumps(inner, ensure_ascii=False) if not isinstance(inner, str) else inner
            return json.dumps(val, ensure_ascii=False)
        if isinstance(val, list):
            return json.dumps(val, ensure_ascii=False)
        return str(val)

    if name in context:
        return _extract(context[name])
    parts = name.split(".", 1)
    if len(parts) == 2:
        node_id, field = parts
        if node_id in context:
            val = context[node_id]
            if isinstance(val, dict) and field in val:
                return _extract(val[field])
    return "{" + name + "}"


def _replace_vars(template: str, context: dict) -> str:
    """替换模板中的 {{变量名}}，支持嵌套字段访问"""
    import re
    def _replacer(m):
        return _resolve_var(m.group(1), context)
    return re.sub(r"\{\{(.+?)\}\}", _replacer, template)


def _execute_node(node: dict, context: dict, llm_router=None) -> Any:
    ntype = node["type"]
    config = node.get("config", {})

    if ntype == "start":
        return context.get("inputs", {})

    elif ntype == "llm":
        if not llm_router:
            return {"error": "LLM router not available"}
        provider = config.get("provider", "agnes")
        prompt = config.get("prompt", "")
        messages = []
        if prompt:
            prompt = _replace_vars(prompt, context)
            messages.append({"role": "user", "content": prompt})
        else:
            user_input = context.get("input", context.get("text", ""))
            messages.append({"role": "user", "content": str(user_input)})

        import asyncio
        import threading

        result_container = [None]
        error_container = [None]

        def _run_in_thread():
            new_loop = asyncio.new_event_loop()
            asyncio.set_event_loop(new_loop)
            try:
                async def _call():
                    result = ""
                    async for chunk in llm_router.chat_stream(messages, provider=provider):
                        if chunk == "[DONE]":
                            continue
                        if chunk.startswith('{"type": "reasoning"'):
                            continue
                        result += chunk
                    return result
                result_container[0] = new_loop.run_until_complete(_call())
            except Exception as e:
                error_container[0] = str(e)
            finally:
                new_loop.close()

        t = threading.Thread(target=_run_in_thread)
        t.start()
        t.join(timeout=120)

        if error_container[0]:
            return {"error": error_container[0]}
        return {"response": result_container[0]}

    elif ntype == "tool":
        tool_name = config.get("tool_name", "")
        tool_input = config.get("input_template", "")
        tool_input = _replace_vars(tool_input, context)

        if tool_name == "ocr":
            try:
                from ocr_engine import ocr_recognize
                import base64 as b64
                ocr_source = config.get("ocr_source", "context")
                if ocr_source == "path":
                    file_path = tool_input.strip()
                    if not file_path or not os.path.exists(file_path):
                        return {"error": f"文件不存在: {file_path}"}
                    with open(file_path, "rb") as f:
                        img_bytes = f.read()
                    return ocr_recognize(img_bytes)
                elif ocr_source == "base64":
                    file_data = tool_input.strip()
                    if not file_data:
                        return {"error": "未提供 base64 图像数据"}
                    if file_data.startswith("data:"):
                        file_data = file_data.split(",", 1)[1]
                    img_bytes = b64.b64decode(file_data)
                    return ocr_recognize(img_bytes)
                else:
                    file_data = context.get("file_data") or context.get("file_base64") or context.get("image_base64", "")
                    if file_data:
                        if file_data.startswith("data:"):
                            file_data = file_data.split(",", 1)[1]
                        img_bytes = b64.b64decode(file_data)
                        return ocr_recognize(img_bytes)
                    elif tool_input:
                        if os.path.exists(tool_input):
                            with open(tool_input, "rb") as f:
                                img_bytes = f.read()
                            return ocr_recognize(img_bytes)
                        return {"error": f"文件不存在: {tool_input}"}
                    return {"error": "OCR 需要提供文件路径或 base64 图像数据"}
            except Exception as e:
                return {"error": f"OCR 失败: {str(e)}"}

        elif tool_name == "http":
            import httpx
            method = config.get("method", "GET").upper()
            url = tool_input or config.get("url", "")
            headers_str = config.get("headers", "")
            body = config.get("body", "")
            timeout_val = int(config.get("timeout", 30))

            url = _replace_vars(url, context)
            body = _replace_vars(body, context)

            headers = {}
            if headers_str:
                try:
                    headers = json.loads(headers_str) if headers_str.strip().startswith("{") else dict(
                        line.split(":", 1) for line in headers_str.strip().split("\n") if ":" in line
                    )
                except Exception:
                    headers = {}

            try:
                with httpx.Client(timeout=timeout_val, follow_redirects=True) as client:
                    kwargs = {"url": url, "headers": headers}
                    if method in ("POST", "PUT", "PATCH"):
                        if body:
                            try:
                                kwargs["json"] = json.loads(body)
                            except Exception:
                                kwargs["data"] = body
                    resp = client.request(method, **kwargs)
                    try:
                        resp_data = resp.json()
                    except Exception:
                        resp_data = resp.text
                    return {
                        "status_code": resp.status_code,
                        "headers": dict(resp.headers),
                        "body": resp_data
                    }
            except httpx.TimeoutException:
                return {"error": f"HTTP 请求超时 ({timeout_val}s)"}
            except Exception as e:
                return {"error": f"HTTP 请求失败: {str(e)}"}

        elif tool_name == "knowledge_search":
            try:
                from langchain_engine import get_rag_engine
                rag = get_rag_engine()
                top_k = int(config.get("top_k", 3))
                results = rag.search(tool_input, top_k=top_k)
                return {"results": results}
            except Exception as e:
                return {"error": str(e)}

        elif tool_name == "text_extract":
            try:
                result = tool_input[:int(config.get("max_length", 5000))]
                return {"result": result}
            except Exception as e:
                return {"error": str(e)}

        elif tool_name == "regex":
            import re
            pattern = config.get("pattern", "")
            text = tool_input
            try:
                matches = re.findall(pattern, text)
                return {"matches": matches, "count": len(matches)}
            except re.error as e:
                return {"error": f"正则表达式错误: {str(e)}"}

        elif tool_name == "json_parse":
            try:
                data = json.loads(tool_input)
                path = config.get("path", "")
                if path:
                    for key in path.split("."):
                        if key.isdigit():
                            data = data[int(key)]
                        else:
                            data = data[key]
                return {"result": data}
            except Exception as e:
                return {"error": f"JSON 解析失败: {str(e)}"}

        elif tool_name == "run_command":
            import subprocess
            import platform
            cmd = tool_input
            if not cmd:
                return {"error": "命令为空"}
            timeout_val = int(config.get("timeout", 30))
            system = platform.system().lower()
            try:
                proc = subprocess.run(cmd, shell=True, capture_output=True, timeout=timeout_val)
                stdout = proc.stdout
                stderr = proc.stderr
                if system == "windows":
                    for enc in ["gbk", "gb2312", "utf-8", "latin-1"]:
                        try:
                            stdout = proc.stdout.decode(enc)
                            stderr = proc.stderr.decode(enc)
                            break
                        except (UnicodeDecodeError, LookupError):
                            continue
                else:
                    stdout = proc.stdout.decode("utf-8", errors="replace")
                    stderr = proc.stderr.decode("utf-8", errors="replace")
                output = stdout[:5000]
                return {
                    "result": output,
                    "stdout": output,
                    "stderr": stderr[:2000],
                    "returncode": proc.returncode
                }
            except subprocess.TimeoutExpired:
                return {"error": f"命令执行超时 ({timeout_val}s)"}
            except Exception as e:
                return {"error": f"命令执行失败: {str(e)}"}

        elif tool_name == "open_app":
            import subprocess
            import platform
            import shutil
            app_name = tool_input.strip()
            if not app_name:
                return {"error": "应用名称为空"}
            system = platform.system().lower()
            app_map = {
                "notepad": "notepad.exe",
                "记事本": "notepad.exe",
                "calculator": "calc.exe",
                "计算器": "calc.exe",
                "paint": "mspaint.exe",
                "画图": "mspaint.exe",
                "explorer": "explorer.exe",
                "文件管理器": "explorer.exe",
                "cmd": "cmd.exe",
                "终端": "cmd.exe",
                "powershell": "powershell.exe",
                "chrome": "chrome",
                "谷歌浏览器": "chrome",
                "edge": "msedge",
                "edge浏览器": "msedge",
                "firefox": "firefox",
                "火狐": "firefox",
                "vscode": "code",
                "vs code": "code",
                "word": "winword",
                "excel": "excel",
                "ppt": "powerpnt",
                "terminal": "wt.exe",
            }
            resolved = app_map.get(app_name.lower(), app_name)
            try:
                if system == "windows":
                    if resolved.endswith(".exe"):
                        subprocess.Popen(f"start {resolved}", shell=True)
                    else:
                        try:
                            subprocess.Popen(["cmd", "/c", "start", "", resolved], shell=False)
                        except FileNotFoundError:
                            found = shutil.which(resolved)
                            if found:
                                subprocess.Popen([found], shell=False)
                            else:
                                subprocess.Popen(f"start {resolved}", shell=True)
                elif system == "darwin":
                    subprocess.Popen(["open", "-a", app_name])
                else:
                    subprocess.Popen([app_name])
                return {"result": f"已启动: {app_name}", "app": app_name, "resolved": resolved}
            except Exception as e:
                return {"error": f"启动应用失败: {str(e)}"}

        elif tool_name == "file_read":
            try:
                file_path = tool_input.strip()
                if not file_path or not os.path.exists(file_path):
                    return {"error": f"文件不存在: {file_path}"}
                encoding = config.get("encoding", "utf-8")
                max_size = int(config.get("max_size", 10000))
                with open(file_path, "r", encoding=encoding, errors="replace") as f:
                    content = f.read(max_size)
                return {"result": content, "path": file_path, "size": os.path.getsize(file_path)}
            except Exception as e:
                return {"error": f"读取文件失败: {str(e)}"}

        elif tool_name == "send_email":
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart

            smtp_host = config.get("smtp_host", "smtp.qq.com")
            smtp_port = int(config.get("smtp_port", 465))
            smtp_user = config.get("smtp_user", "")
            smtp_pass = config.get("smtp_pass", "")
            use_ssl = config.get("use_ssl", True)

            if not smtp_user or not smtp_pass:
                return {"error": "未配置 SMTP 账号或密码"}

            to_addr = config.get("to_addr", "")
            subject = config.get("subject", "")
            body = tool_input
            try:
                parsed = json.loads(body)
                if isinstance(parsed, dict):
                    if "result" in parsed:
                        body = str(parsed["result"])
                    elif "response" in parsed:
                        body = str(parsed["response"])
                    elif "stdout" in parsed:
                        body = str(parsed["stdout"])
            except (json.JSONDecodeError, TypeError):
                pass
            body = body.encode('utf-8', errors='replace').decode('utf-8', errors='replace')

            if not to_addr:
                return {"error": "未配置收件人"}

            to_addr = _replace_vars(to_addr, context)
            subject = _replace_vars(subject, context)

            msg = MIMEMultipart()
            msg["From"] = smtp_user
            msg["To"] = to_addr
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "plain", "utf-8"))

            try:
                if use_ssl:
                    server = smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=15)
                else:
                    server = smtplib.SMTP(smtp_host, smtp_port, timeout=15)
                    server.starttls()
                server.login(smtp_user, smtp_pass)
                server.sendmail(smtp_user, to_addr.split(","), msg.as_string())
                server.quit()
                return {"result": f"邮件已发送", "to": to_addr, "subject": subject}
            except smtplib.SMTPAuthenticationError:
                return {"error": "SMTP 认证失败，请检查账号密码"}
            except smtplib.SMTPConnectError:
                return {"error": f"无法连接 SMTP 服务器 {smtp_host}:{smtp_port}"}
            except Exception as e:
                return {"error": f"发送邮件失败: {str(e)}"}

        elif tool_name == "file_write":
            try:
                file_path = config.get("file_path", "")
                file_path = _replace_vars(file_path, context)
                if not file_path:
                    return {"error": "未指定写入路径"}
                write_mode = config.get("write_mode", "overwrite")
                content = tool_input
                os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)
                if write_mode == "append" and os.path.exists(file_path):
                    with open(file_path, "a", encoding="utf-8", errors="replace") as f:
                        f.write(content)
                else:
                    with open(file_path, "w", encoding="utf-8", errors="replace") as f:
                        f.write(content)
                return {"result": f"已写入文件", "path": file_path, "size": len(content.encode("utf-8")), "mode": write_mode}
            except Exception as e:
                return {"error": f"写入文件失败: {str(e)}"}

        elif tool_name == "pdf_extract":
            try:
                pdf_path = tool_input.strip()
                if not pdf_path or not os.path.exists(pdf_path):
                    return {"error": f"PDF 文件不存在: {pdf_path}"}
                try:
                    import fitz
                    doc = fitz.open(pdf_path)
                    pages = []
                    full_text = []
                    for i, page in enumerate(doc):
                        text = page.get_text()
                        pages.append({"page": i + 1, "text": text})
                        full_text.append(text)
                    doc.close()
                    return {"result": "\n".join(full_text), "pages": pages, "total_pages": len(pages)}
                except ImportError:
                    import subprocess
                    result = subprocess.run(
                        ["python", "-m", "PyPDF2"],
                        capture_output=True, timeout=5
                    )
                    raise ImportError("需要安装 PyMuPDF: pip install PyMuPDF")
            except ImportError as e:
                return {"error": f"PDF 解析库未安装: {str(e)}。请运行: pip install PyMuPDF"}
            except Exception as e:
                return {"error": f"PDF 解析失败: {str(e)}"}

        elif tool_name == "clipboard":
            try:
                import subprocess
                platform_name = __import__("platform").system().lower()
                action = config.get("action", "paste")
                if action == "paste":
                    if platform_name == "windows":
                        result = subprocess.run(["powershell", "-command", "Get-Clipboard"],
                                               capture_output=True, timeout=5, shell=False)
                        text = result.stdout.decode("utf-8", errors="replace").strip()
                    elif platform_name == "darwin":
                        result = subprocess.run(["pbpaste"], capture_output=True, timeout=5)
                        text = result.stdout.decode("utf-8", errors="replace").strip()
                    else:
                        result = subprocess.run(["xclip", "-selection", "clipboard", "-o"],
                                               capture_output=True, timeout=5)
                        text = result.stdout.decode("utf-8", errors="replace").strip()
                    return {"result": text}
                elif action == "copy":
                    text = tool_input
                    if platform_name == "windows":
                        subprocess.run(["powershell", "-command", f"Set-Clipboard -Value '{text}'"],
                                       timeout=5, shell=False)
                    elif platform_name == "darwin":
                        proc = subprocess.Popen(["pbcopy"], stdin=subprocess.PIPE)
                        proc.communicate(text.encode("utf-8"))
                    else:
                        proc = subprocess.Popen(["xclip", "-selection", "clipboard"], stdin=subprocess.PIPE)
                        proc.communicate(text.encode("utf-8"))
                    return {"result": "已复制到剪贴板"}
                else:
                    return {"error": f"未知剪贴板操作: {action}"}
            except Exception as e:
                return {"error": f"剪贴板操作失败: {str(e)}"}

        elif tool_name == "screenshot":
            try:
                save_path = config.get("save_path", "")
                save_path = _replace_vars(save_path, context)
                if not save_path:
                    save_path = os.path.join("generated_media", f"screenshot_{int(__import__('time').time())}.png")
                os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
                try:
                    import pyautogui
                    img = pyautogui.screenshot()
                    img.save(save_path)
                except ImportError:
                    import subprocess
                    if __import__("platform").system().lower() == "windows":
                        ps_cmd = (
                            "Add-Type -AssemblyName System.Windows.Forms; "
                            "$screen = [System.Windows.Forms.Screen]::PrimaryScreen; "
                            "$bmp = New-Object System.Drawing.Bitmap($screen.Bounds.Width, $screen.Bounds.Height); "
                            "$gfx = [System.Drawing.Graphics]::FromImage($bmp); "
                            "$gfx.CopyFromScreen($screen.Bounds.Location, [System.Drawing.Point]::Empty, $screen.Bounds.Size); "
                            f"$bmp.Save('{save_path}')"
                        )
                        subprocess.run(["powershell", "-command", ps_cmd], timeout=10, shell=False)
                    else:
                        return {"error": "截图需要安装 pyautogui: pip install pyautogui"}
                return {"result": f"截图已保存", "path": save_path}
            except Exception as e:
                return {"error": f"截图失败: {str(e)}"}

    elif ntype == "condition":
        condition = config.get("condition", "true")
        try:
            condition = _replace_vars(condition, context)
            result = eval(condition)
            return {"result": bool(result), "branch": "true" if result else "false"}
        except Exception:
            return {"result": False, "branch": "false"}

    elif ntype == "transform":
        template = config.get("template", "{{input}}")
        template = _replace_vars(template, context)
        return {"result": template}

    elif ntype == "delay":
        import time
        seconds = int(config.get("seconds", 3))
        seconds = min(seconds, 60)
        time.sleep(seconds)
        return {"result": f"waited {seconds} seconds"}

    elif ntype == "loop":
        list_var = config.get("list_var", "input")
        raw = context.get(list_var, context.get("inputs", {}).get(list_var, []))
        if isinstance(raw, str):
            try:
                raw = json.loads(raw)
            except Exception:
                raw = [raw]
        if not isinstance(raw, list):
            raw = [raw]
        return {"items": raw, "count": len(raw), "current": raw[0] if raw else None}

    elif ntype == "parallel":
        upstream_results = {}
        current_node_id = node.get("id", "")
        for k, v in context.items():
            if k not in ("inputs",) and k != current_node_id:
                upstream_results[k] = v
        return {"results": upstream_results, "count": len(upstream_results)}

    elif ntype == "database":
        db_type = config.get("db_type", "mongodb")
        connection_string = _replace_vars(config.get("connection_string", ""), context)
        query = _replace_vars(config.get("query", ""), context)
        if not connection_string:
            return {"error": "未配置数据库连接字符串"}
        if not query:
            return {"error": "未配置查询语句"}
        try:
            if db_type == "mongodb":
                from pymongo import MongoClient
                client = MongoClient(connection_string, serverSelectionTimeoutMS=5000)
                db_name = config.get("db_name", "test")
                coll_name = config.get("collection", "test")
                coll = client[db_name][coll_name]
                query_obj = json.loads(query) if query.strip().startswith("{") else {}
                results = list(coll.find(query_obj).limit(100))
                for r in results:
                    r["_id"] = str(r["_id"])
                return {"result": results, "count": len(results)}
            elif db_type in ("mysql", "postgresql"):
                if db_type == "mysql":
                    import pymysql
                    conn = pymysql.connect(connection_string, connect_timeout=10)
                else:
                    import psycopg2
                    conn = psycopg2.connect(connection_string, connect_timeout=10)
                cursor = conn.cursor()
                cursor.execute(query)
                if cursor.description:
                    columns = [d[0] for d in cursor.description]
                    rows = [dict(zip(columns, row)) for row in cursor.fetchall()]
                    return {"result": rows, "count": len(rows)}
                else:
                    conn.commit()
                    return {"result": f"affected {cursor.rowcount} rows", "count": cursor.rowcount}
            else:
                return {"error": f"不支持的数据库类型: {db_type}"}
        except Exception as e:
            return {"error": f"数据库操作失败: {str(e)}"}

    elif ntype == "file_operation":
        operation = config.get("operation", "read")
        file_path = _replace_vars(config.get("path", ""), context)
        if not file_path:
            return {"error": "未配置文件路径"}
        try:
            if operation == "read":
                encoding = config.get("encoding", "utf-8")
                with open(file_path, "r", encoding=encoding, errors="replace") as f:
                    content = f.read(int(config.get("max_size", 50000)))
                return {"result": content, "path": file_path, "size": os.path.getsize(file_path)}
            elif operation == "write":
                content = _replace_vars(config.get("content", ""), context)
                os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(content)
                return {"result": "写入成功", "path": file_path}
            elif operation == "copy":
                dest = _replace_vars(config.get("dest", ""), context)
                import shutil
                shutil.copy2(file_path, dest)
                return {"result": "复制成功", "src": file_path, "dest": dest}
            elif operation == "move":
                dest = _replace_vars(config.get("dest", ""), context)
                import shutil
                shutil.move(file_path, dest)
                return {"result": "移动成功", "src": file_path, "dest": dest}
            elif operation == "delete":
                if os.path.isdir(file_path):
                    import shutil
                    shutil.rmtree(file_path)
                else:
                    os.remove(file_path)
                return {"result": "删除成功", "path": file_path}
            elif operation == "list":
                items = os.listdir(file_path) if os.path.isdir(file_path) else [file_path]
                return {"result": items, "count": len(items), "path": file_path}
            else:
                return {"error": f"未知操作: {operation}"}
        except Exception as e:
            return {"error": f"文件操作失败: {str(e)}"}

    elif ntype == "webhook":
        import httpx
        method = config.get("method", "GET").upper()
        url = _replace_vars(config.get("url", ""), context)
        headers_str = config.get("headers", "")
        body = _replace_vars(config.get("body", ""), context)
        timeout_val = int(config.get("timeout", 30))
        if not url:
            return {"error": "未配置 URL"}
        headers = {}
        if headers_str:
            try:
                headers = json.loads(headers_str) if headers_str.strip().startswith("{") else dict(
                    line.split(":", 1) for line in headers_str.strip().split("\n") if ":" in line
                )
            except Exception:
                headers = {}
        try:
            with httpx.Client(timeout=timeout_val, follow_redirects=True) as client:
                kwargs = {"url": url, "headers": headers}
                if method in ("POST", "PUT", "PATCH"):
                    if body:
                        try:
                            kwargs["json"] = json.loads(body)
                        except Exception:
                            kwargs["data"] = body
                resp = client.request(method, **kwargs)
                try:
                    resp_data = resp.json()
                except Exception:
                    resp_data = resp.text
                return {"status_code": resp.status_code, "body": resp_data, "headers": dict(resp.headers)}
        except httpx.TimeoutException:
            return {"error": f"请求超时 ({timeout_val}s)"}
        except Exception as e:
            return {"error": f"请求失败: {str(e)}"}

    elif ntype == "code_exec":
        language = config.get("language", "python")
        code = _replace_vars(config.get("code", ""), context)
        timeout_val = int(config.get("timeout", 30))
        if not code:
            return {"error": "代码为空"}
        import subprocess
        import tempfile
        try:
            if language == "python":
                with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False, encoding="utf-8") as f:
                    f.write(code)
                    tmp_path = f.name
                proc = subprocess.run(
                    ["python", tmp_path], capture_output=True, timeout=timeout_val,
                    env={**os.environ, "PYTHONIOENCODING": "utf-8"}
                )
            elif language in ("javascript", "js"):
                with tempfile.NamedTemporaryFile(mode="w", suffix=".js", delete=False, encoding="utf-8") as f:
                    f.write(code)
                    tmp_path = f.name
                proc = subprocess.run(["node", tmp_path], capture_output=True, timeout=timeout_val)
            else:
                return {"error": f"不支持的语言: {language}"}
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            stdout = proc.stdout.decode("utf-8", errors="replace")[:5000]
            stderr = proc.stderr.decode("utf-8", errors="replace")[:2000]
            return {"result": stdout, "stderr": stderr, "returncode": proc.returncode}
        except subprocess.TimeoutExpired:
            return {"error": f"代码执行超时 ({timeout_val}s)"}
        except Exception as e:
            return {"error": f"执行失败: {str(e)}"}

    elif ntype == "image_gen":
        prompt = _replace_vars(config.get("prompt", ""), context)
        if not prompt:
            return {"error": "未配置图像 Prompt"}
        import httpx as _httpx
        size = config.get("size", "1024x1024")
        try:
            from pymongo import MongoClient as _MC
            _c = _MC(os.environ.get("MONGODB_URL", "mongodb://ai_mongo:ai_mongo_2024@localhost:27017"), serverSelectionTimeoutMS=3000)
            _s = _c["ai_framework"]["settings"].find_one({"_id": "app_settings"}) or {}
            api_key = _s.get("agnes_key", "")
            base_url = _s.get("agnes_url", "https://apihub.agnes-ai.com")
            model = config.get("model") or _s.get("agnes_image_model_1", "agnes-image-2.1-flash")
            if not api_key:
                return {"error": "未配置图像生成 API Key"}
            url = f"{base_url.rstrip('/')}/v1/images/generations"
            payload = {"model": model, "prompt": prompt, "n": 1, "size": size}
            with _httpx.Client(timeout=120.0) as client:
                resp = client.post(url, headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}, json=payload)
                resp.raise_for_status()
                data = resp.json()
            import base64
            media_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_media", "images")
            os.makedirs(media_dir, exist_ok=True)
            saved = []
            for i, img in enumerate(data.get("data", [])):
                b64 = img.get("b64_json")
                img_url = img.get("url")
                if b64:
                    filename = f"wf_img_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{i}.png"
                    filepath = os.path.join(media_dir, filename)
                    with open(filepath, "wb") as f:
                        f.write(base64.b64decode(b64))
                    saved.append({"filename": filename, "path": filepath})
                elif img_url:
                    saved.append({"url": img_url})
            return {"result": saved[0] if saved else "生成完成", "images": saved}
        except Exception as e:
            return {"error": f"图像生成失败: {str(e)}"}

    elif ntype == "error_handler":
        fallback = config.get("fallback_output", "上游执行出错")
        upstream_error = None
        for k, v in context.items():
            if isinstance(v, dict) and "error" in v:
                upstream_error = v["error"]
                break
        if upstream_error:
            return {"result": fallback, "original_error": upstream_error}
        return {"result": "上游执行正常，无需处理"}

    elif ntype == "text_split":
        delimiter = config.get("delimiter", "\n")
        text = _replace_vars(config.get("text", "{{input}}"), context)
        parts = [p.strip() for p in text.split(delimiter) if p.strip()]
        return {"items": parts, "count": len(parts)}

    elif ntype == "text_translate":
        if not llm_router:
            return {"error": "LLM router not available"}
        text = _replace_vars(config.get("text", "{{input}}"), context)
        target_lang = config.get("target_lang", "英语")
        prompt = f"将以下文本翻译为{target_lang}，只输出翻译结果：\n{text}"
        import asyncio, threading
        result_container = [None]
        error_container = [None]
        def _run():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                async def _call():
                    r = ""
                    async for chunk in llm_router.chat_stream([{"role": "user", "content": prompt}], provider=config.get("provider", "agnes")):
                        if chunk == "[DONE]" or chunk.startswith('{"type": "reasoning"'):
                            continue
                        r += chunk
                    return r
                result_container[0] = loop.run_until_complete(_call())
            except Exception as e:
                error_container[0] = str(e)
            finally:
                loop.close()
        t = threading.Thread(target=_run); t.start(); t.join(timeout=120)
        if error_container[0]:
            return {"error": error_container[0]}
        return {"result": result_container[0], "source": text, "target_lang": target_lang}

    elif ntype == "text_summarize":
        if not llm_router:
            return {"error": "LLM router not available"}
        text = _replace_vars(config.get("text", "{{input}}"), context)
        max_length = config.get("max_length", 200)
        prompt = f"用中文总结以下内容，不超过{max_length}字：\n{text[:3000]}"
        import asyncio, threading
        result_container = [None]
        error_container = [None]
        def _run():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                async def _call():
                    r = ""
                    async for chunk in llm_router.chat_stream([{"role": "user", "content": prompt}], provider=config.get("provider", "agnes")):
                        if chunk == "[DONE]" or chunk.startswith('{"type": "reasoning"'):
                            continue
                        r += chunk
                    return r
                result_container[0] = loop.run_until_complete(_call())
            except Exception as e:
                error_container[0] = str(e)
            finally:
                loop.close()
        t = threading.Thread(target=_run); t.start(); t.join(timeout=120)
        if error_container[0]:
            return {"error": error_container[0]}
        return {"result": result_container[0], "original_length": len(text)}

    elif ntype == "json_build":
        fields_str = config.get("fields", "{}")
        fields_str = _replace_vars(fields_str, context)
        try:
            fields = json.loads(fields_str)
        except Exception:
            fields = {}
        key_value_pairs = config.get("key_value_pairs", "")
        if key_value_pairs:
            for line in key_value_pairs.strip().split("\n"):
                if "=" in line:
                    k, v = line.split("=", 1)
                    k = k.strip()
                    v = v.strip()
                    v_resolved = _replace_vars(v, context)
                    try:
                        fields[k] = json.loads(v_resolved)
                    except Exception:
                        fields[k] = v_resolved
        return {"result": fields}

    elif ntype == "data_filter":
        list_var = config.get("list_var", "input")
        raw = context.get(list_var, context.get("inputs", {}).get(list_var, []))
        if isinstance(raw, str):
            try: raw = json.loads(raw)
            except: raw = [raw]
        if not isinstance(raw, list):
            raw = [raw]
        field = config.get("field", "")
        op = config.get("operator", "eq")
        value = _replace_vars(config.get("value", ""), context)
        filtered = []
        for item in raw:
            if isinstance(item, dict) and field:
                item_val = item.get(field, "")
            else:
                item_val = item
            try:
                if op == "eq" and str(item_val) == value: filtered.append(item)
                elif op == "ne" and str(item_val) != value: filtered.append(item)
                elif op == "gt" and float(item_val) > float(value): filtered.append(item)
                elif op == "lt" and float(item_val) < float(value): filtered.append(item)
                elif op == "contains" and value in str(item_val): filtered.append(item)
                elif op == "startswith" and str(item_val).startswith(value): filtered.append(item)
            except (ValueError, TypeError):
                continue
        return {"items": filtered, "count": len(filtered), "original_count": len(raw)}

    elif ntype == "data_sort":
        list_var = config.get("list_var", "input")
        raw = context.get(list_var, context.get("inputs", {}).get(list_var, []))
        if isinstance(raw, str):
            try: raw = json.loads(raw)
            except: raw = [raw]
        if not isinstance(raw, list):
            raw = [raw]
        sort_field = config.get("sort_field", "")
        reverse = config.get("reverse", False)
        if sort_field and raw and isinstance(raw[0], dict):
            sorted_list = sorted(raw, key=lambda x: x.get(sort_field, ""), reverse=reverse)
        else:
            sorted_list = sorted(raw, reverse=reverse)
        return {"items": sorted_list, "count": len(sorted_list)}

    elif ntype == "switch":
        switch_var = _replace_vars(config.get("switch_var", "{{input}}"), context)
        cases_str = config.get("cases", "")
        default_case = config.get("default_case", "default")
        matched = default_case
        for line in cases_str.strip().split("\n"):
            if ":" in line:
                case_val, case_name = line.split(":", 1)
                if case_val.strip() == str(switch_var).strip():
                    matched = case_name.strip()
                    break
        return {"result": switch_var, "matched": matched, "branch": matched}

    elif ntype == "sub_workflow":
        sub_wf_id = config.get("workflow_id", "")
        if not sub_wf_id:
            return {"error": "未配置子工作流 ID"}
        sub_inputs = {}
        inputs_str = config.get("inputs_map", "{}")
        inputs_str = _replace_vars(inputs_str, context)
        try:
            sub_inputs = json.loads(inputs_str)
        except Exception:
            sub_inputs = {"input": str(context.get("input", ""))}
        try:
            sub_run = run_workflow(sub_wf_id, inputs=sub_inputs, llm_router=llm_router)
            return {"result": sub_run.outputs, "status": sub_run.status, "workflow_id": sub_wf_id}
        except Exception as e:
            return {"error": f"子工作流执行失败: {str(e)}"}

    elif ntype == "retry":
        max_retries = int(config.get("max_retries", 3))
        delay_sec = int(config.get("delay", 2))
        import time
        last_error = None
        for attempt in range(max_retries):
            upstream_ok = True
            for k, v in context.items():
                if isinstance(v, dict) and "error" in v:
                    upstream_ok = False
                    last_error = v["error"]
                    break
            if upstream_ok:
                return {"result": "上游执行成功", "attempts": attempt + 1}
            if attempt < max_retries - 1:
                time.sleep(delay_sec)
        return {"error": f"重试{max_retries}次后仍然失败: {last_error}", "attempts": max_retries}

    elif ntype == "notify":
        notify_type = config.get("notify_type", "dingtalk")
        webhook_url = _replace_vars(config.get("webhook_url", ""), context)
        content = _replace_vars(config.get("content", "{{input}}"), context)
        title = _replace_vars(config.get("title", "工作流通知"), context)
        if not webhook_url:
            return {"error": "未配置 Webhook URL"}
        import httpx as _httpx
        try:
            if notify_type == "dingtalk":
                payload = {"msgtype": "text", "text": {"content": f"{title}\n\n{content}"}}
            elif notify_type == "wecom":
                payload = {"msgtype": "text", "text": {"content": f"{title}\n\n{content}"}}
            elif notify_type == "feishu":
                payload = {"msg_type": "text", "content": {"text": f"{title}\n\n{content}"}}
            else:
                payload = {"text": f"{title}\n\n{content}"}
            with _httpx.Client(timeout=10) as client:
                resp = client.post(webhook_url, json=payload)
                return {"result": f"通知已发送 ({notify_type})", "status_code": resp.status_code}
        except Exception as e:
            return {"error": f"发送通知失败: {str(e)}"}

    # ==================== 数据处理类 ====================

    elif ntype == "math_calc":
        expression = config.get("expression", "0")
        expression = _replace_vars(expression, context)
        try:
            import math
            safe_dict = {
                "abs": abs, "round": round, "min": min, "max": max,
                "sum": sum, "pow": pow, "int": int, "float": float,
                "sqrt": math.sqrt, "sin": math.sin, "cos": math.cos,
                "tan": math.tan, "log": math.log, "log10": math.log10,
                "pi": math.pi, "e": math.e, "ceil": math.ceil, "floor": math.floor,
            }
            result = eval(expression, {"__builtins__": {}}, safe_dict)
            return {"result": result, "expression": expression}
        except Exception as e:
            return {"error": f"数学计算失败: {str(e)}"}

    elif ntype == "datetime":
        action = config.get("action", "now")
        fmt = config.get("format", "%Y-%m-%d %H:%M:%S")
        import datetime as _dt
        try:
            if action == "now":
                result = _dt.datetime.now().strftime(fmt)
            elif action == "today":
                result = _dt.date.today().strftime(fmt)
            elif action == "timestamp":
                result = int(_dt.datetime.now().timestamp())
            elif action == "parse":
                text = _replace_vars(config.get("input_template", ""), context)
                result = _dt.datetime.strptime(text, fmt).isoformat()
            elif action == "add":
                text = _replace_vars(config.get("input_template", ""), context)
                days = int(config.get("days", 0))
                hours = int(config.get("hours", 0))
                minutes = int(config.get("minutes", 0))
                dt = _dt.datetime.fromisoformat(text) if text else _dt.datetime.now()
                result = (dt + _dt.timedelta(days=days, hours=hours, minutes=minutes)).strftime(fmt)
            elif action == "diff":
                text1 = _replace_vars(config.get("input_template", ""), context)
                text2 = _replace_vars(config.get("end_date", ""), context)
                dt1 = _dt.datetime.fromisoformat(text1)
                dt2 = _dt.datetime.fromisoformat(text2)
                diff = dt2 - dt1
                result = {"days": diff.days, "seconds": diff.seconds, "total_hours": diff.total_seconds() / 3600}
            else:
                result = _dt.datetime.now().strftime(fmt)
            return {"result": result, "action": action}
        except Exception as e:
            return {"error": f"日期时间操作失败: {str(e)}"}

    elif ntype == "type_convert":
        target_type = config.get("target_type", "string")
        input_val = _replace_vars(config.get("input_template", ""), context)
        try:
            if target_type == "string":
                result = str(input_val)
            elif target_type == "int":
                result = int(float(input_val))
            elif target_type == "float":
                result = float(input_val)
            elif target_type == "bool":
                result = bool(input_val) and input_val not in ("0", "false", "False", "null", "None", "")
            elif target_type == "list":
                if isinstance(input_val, str):
                    result = [x.strip() for x in input_val.split(config.get("delimiter", ","))]
                elif isinstance(input_val, list):
                    result = input_val
                else:
                    result = [input_val]
            elif target_type == "json":
                result = json.loads(input_val) if isinstance(input_val, str) else json.loads(json.dumps(input_val))
            elif target_type == "dict":
                if isinstance(input_val, str):
                    result = dict(item.split("=", 1) for item in input_val.split("\n") if "=" in item)
                elif isinstance(input_val, dict):
                    result = input_val
                else:
                    result = {"value": input_val}
            else:
                result = str(input_val)
            return {"result": result, "type": target_type}
        except Exception as e:
            return {"error": f"类型转换失败: {str(e)}"}

    elif ntype == "csv_parse":
        delimiter = config.get("delimiter", ",")
        has_header = config.get("has_header", True)
        input_val = _replace_vars(config.get("input_template", ""), context)
        try:
            import csv
            import io
            if os.path.isfile(input_val):
                with open(input_val, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()
            else:
                content = input_val
            reader = csv.DictReader(io.StringIO(content), delimiter=delimiter) if has_header else csv.reader(io.StringIO(content), delimiter=delimiter)
            if has_header:
                rows = [dict(row) for row in reader]
            else:
                rows = [row for row in reader]
            return {"result": rows, "count": len(rows)}
        except Exception as e:
            return {"error": f"CSV 解析失败: {str(e)}"}

    elif ntype == "excel_read":
        file_path = _replace_vars(config.get("input_template", ""), context)
        sheet_name = config.get("sheet_name", "")
        max_rows = int(config.get("max_rows", 1000))
        try:
            import openpyxl
            if not file_path or not os.path.exists(file_path):
                return {"error": f"Excel 文件不存在: {file_path}"}
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows:
                    break
                rows.append([str(cell) if cell is not None else "" for cell in row])
            wb.close()
            headers = rows[0] if rows else []
            data = [dict(zip(headers, row)) for row in rows[1:]] if headers else rows
            return {"result": data, "headers": headers, "count": len(data)}
        except ImportError:
            return {"error": "需要安装 openpyxl: pip install openpyxl"}
        except Exception as e:
            return {"error": f"Excel 读取失败: {str(e)}"}

    elif ntype == "regex_replace":
        pattern = config.get("pattern", "")
        replacement = config.get("replacement", "")
        input_val = _replace_vars(config.get("input_template", ""), context)
        try:
            import re
            result = re.sub(pattern, replacement, input_val)
            count = len(re.findall(pattern, input_val))
            return {"result": result, "count": count}
        except re.error as e:
            return {"error": f"正则表达式错误: {str(e)}"}
        except Exception as e:
            return {"error": f"正则替换失败: {str(e)}"}

    elif ntype == "hash_encode":
        algorithm = config.get("algorithm", "md5")
        encoding = config.get("encoding", "hex")
        input_val = _replace_vars(config.get("input_template", ""), context)
        try:
            import hashlib
            import base64
            data = input_val.encode("utf-8")
            if algorithm == "md5":
                h = hashlib.md5(data)
            elif algorithm == "sha1":
                h = hashlib.sha1(data)
            elif algorithm == "sha256":
                h = hashlib.sha256(data)
            elif algorithm == "sha512":
                h = hashlib.sha512(data)
            elif algorithm == "base64_encode":
                result = base64.b64encode(data).decode("utf-8")
                return {"result": result, "algorithm": algorithm}
            elif algorithm == "base64_decode":
                result = base64.b64decode(input_val).decode("utf-8", errors="replace")
                return {"result": result, "algorithm": algorithm}
            elif algorithm == "url_encode":
                from urllib.parse import quote
                result = quote(input_val)
                return {"result": result, "algorithm": algorithm}
            elif algorithm == "url_decode":
                from urllib.parse import unquote
                result = unquote(input_val)
                return {"result": result, "algorithm": algorithm}
            else:
                return {"error": f"不支持的算法: {algorithm}"}
            if encoding == "hex":
                result = h.hexdigest()
            elif encoding == "base64":
                result = base64.b64encode(h.digest()).decode("utf-8")
            else:
                result = h.hexdigest()
            return {"result": result, "algorithm": algorithm}
        except Exception as e:
            return {"error": f"哈希/编码失败: {str(e)}"}

    elif ntype == "uuid_generate":
        version = config.get("version", "4")
        count = int(config.get("count", 1))
        try:
            import uuid
            results = []
            for _ in range(min(count, 100)):
                if version == "1":
                    results.append(str(uuid.uuid1()))
                else:
                    results.append(str(uuid.uuid4()))
            return {"result": results[0] if count == 1 else results, "count": len(results)}
        except Exception as e:
            return {"error": f"UUID 生成失败: {str(e)}"}

    # ==================== AI 扩展类 ====================

    elif ntype == "text_embedding":
        input_val = _replace_vars(config.get("input_template", ""), context)
        provider = config.get("provider", "agnes")
        try:
            import httpx as _httpx
            from dotenv import load_dotenv
            load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))
            if provider == "agnes":
                api_key = os.getenv("AGNES_API_KEY", "")
                base_url = os.getenv("AGNES_BASE_URL", "https://apihub.agnes-ai.com")
            else:
                api_key = os.getenv("OPENAI_API_KEY", "")
                base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
            if not api_key:
                return {"error": "未配置 API Key"}
            model = config.get("model", "text-embedding-ada-002")
            texts = [input_val] if isinstance(input_val, str) else input_val
            with _httpx.Client(timeout=30) as client:
                resp = client.post(
                    f"{base_url}/v1/embeddings",
                    headers={"Authorization": f"Bearer {api_key}"},
                    json={"model": model, "input": texts}
                )
                resp.raise_for_status()
                data = resp.json()
                embeddings = [item["embedding"] for item in data["data"]]
                return {"result": embeddings[0] if len(embeddings) == 1 else embeddings, "dimensions": len(embeddings[0])}
        except Exception as e:
            return {"error": f"文本嵌入失败: {str(e)}"}

    elif ntype == "speech_to_text":
        audio_path = _replace_vars(config.get("input_template", ""), context)
        language = config.get("language", "zh")
        try:
            import whisper
            model = whisper.load_model("base")
            result = model.transcribe(audio_path, language=language)
            return {"result": result["text"], "language": result.get("language", language)}
        except ImportError:
            return {"error": "需要安装 openai-whisper: pip install openai-whisper"}
        except Exception as e:
            return {"error": f"语音识别失败: {str(e)}"}

    elif ntype == "text_to_speech":
        input_val = _replace_vars(config.get("input_template", ""), context)
        save_path = config.get("save_path", "")
        voice = config.get("voice", "alloy")
        try:
            from dotenv import load_dotenv
            load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))
            api_key = os.getenv("OPENAI_API_KEY", "")
            base_url = os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1")
            if not api_key:
                return {"error": "未配置 OPENAI_API_KEY"}
            if not save_path:
                save_path = os.path.join("generated_media", f"tts_{int(__import__('time').time())}.mp3")
            os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
            import httpx as _httpx
            with _httpx.Client(timeout=60) as client:
                resp = client.post(
                    f"{base_url}/v1/audio/speech",
                    headers={"Authorization": f"Bearer {api_key}"},
                    json={"model": "tts-1", "input": input_val, "voice": voice}
                )
                resp.raise_for_status()
                with open(save_path, "wb") as f:
                    f.write(resp.content)
            return {"result": f"语音已生成", "path": save_path, "voice": voice}
        except Exception as e:
            return {"error": f"语音合成失败: {str(e)}"}

    elif ntype == "image_process":
        input_path = _replace_vars(config.get("input_template", ""), context)
        action = config.get("action", "resize")
        save_path = config.get("save_path", "")
        try:
            from PIL import Image
            if not input_path or not os.path.exists(input_path):
                return {"error": f"图片不存在: {input_path}"}
            img = Image.open(input_path)
            if action == "resize":
                width = int(config.get("width", 800))
                height = int(config.get("height", 600))
                img = img.resize((width, height), Image.LANCZOS)
            elif action == "crop":
                left = int(config.get("left", 0))
                top = int(config.get("top", 0))
                right = int(config.get("right", img.width))
                bottom = int(config.get("bottom", img.height))
                img = img.crop((left, top, right, bottom))
            elif action == "rotate":
                angle = int(config.get("angle", 90))
                img = img.rotate(angle, expand=True)
            elif action == "watermark":
                from PIL import ImageDraw, ImageFont
                text = config.get("text", "Watermark")
                draw = ImageDraw.Draw(img)
                try:
                    font = ImageFont.truetype("arial.ttf", 36)
                except Exception:
                    font = ImageFont.load_default()
                draw.text((10, img.height - 50), text, fill="rgba(255,255,255,128)", font=font)
            elif action == "grayscale":
                img = img.convert("L")
            elif action == "thumbnail":
                max_size = int(config.get("max_size", 200))
                img.thumbnail((max_size, max_size), Image.LANCZOS)
            if not save_path:
                base, ext = os.path.splitext(input_path)
                save_path = f"{base}_{action}{ext}"
            os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
            img.save(save_path)
            return {"result": f"图片已处理", "path": save_path, "action": action, "size": f"{img.width}x{img.height}"}
        except ImportError:
            return {"error": "需要安装 Pillow: pip install Pillow"}
        except Exception as e:
            return {"error": f"图片处理失败: {str(e)}"}

    elif ntype == "markdown_html":
        input_val = _replace_vars(config.get("input_template", ""), context)
        try:
            import markdown
            html = markdown.markdown(input_val, extensions=["extra", "codehilite", "tables"])
            save_path = config.get("save_path", "")
            if save_path:
                save_path = _replace_vars(save_path, context)
                os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(f"<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>{html}</body></html>")
            return {"result": html, "path": save_path or None}
        except ImportError:
            return {"error": "需要安装 markdown: pip install markdown"}
        except Exception as e:
            return {"error": f"Markdown 转换失败: {str(e)}"}

    # ==================== 网络通信类 ====================

    elif ntype == "websocket_connect":
        url = _replace_vars(config.get("url", ""), context)
        message = _replace_vars(config.get("input_template", ""), context)
        timeout_val = int(config.get("timeout", 10))
        try:
            import websocket
            ws = websocket.create_connection(url, timeout=timeout_val)
            if message:
                ws.send(message)
            result_data = ws.recv()
            ws.close()
            return {"result": result_data, "url": url}
        except ImportError:
            return {"error": "需要安装 websocket-client: pip install websocket-client"}
        except Exception as e:
            return {"error": f"WebSocket 连接失败: {str(e)}"}

    elif ntype == "http_stream":
        url = _replace_vars(config.get("url", ""), context)
        method = config.get("method", "GET").upper()
        headers_str = config.get("headers", "")
        body = config.get("body", "")
        timeout_val = int(config.get("timeout", 30))
        headers = {}
        if headers_str:
            try:
                headers = json.loads(headers_str) if headers_str.strip().startswith("{") else dict(
                    line.split(":", 1) for line in headers_str.strip().split("\n") if ":" in line
                )
            except Exception:
                pass
        body = _replace_vars(body, context)
        try:
            import httpx as _httpx
            collected = []
            with _httpx.Client(timeout=timeout_val, follow_redirects=True) as client:
                kwargs = {"url": url, "headers": headers}
                if method in ("POST", "PUT", "PATCH") and body:
                    try:
                        kwargs["json"] = json.loads(body)
                    except Exception:
                        kwargs["data"] = body
                with client.stream(method, **kwargs) as resp:
                    for line in resp.iter_lines():
                        if line:
                            collected.append(line)
                            if len(collected) >= 500:
                                break
            return {"result": "\n".join(collected), "url": url, "lines": len(collected)}
        except Exception as e:
            return {"error": f"HTTP 流式请求失败: {str(e)}"}

    elif ntype == "network_ping":
        target = _replace_vars(config.get("input_template", ""), context)
        count = int(config.get("count", 4))
        timeout_val = int(config.get("timeout", 5))
        try:
            import subprocess
            param = "-n" if __import__("platform").system().lower() == "windows" else "-c"
            proc = subprocess.run(
                ["ping", param, str(min(count, 10)), "-w", str(timeout_val * 1000), target],
                capture_output=True, timeout=timeout_val * 10 + 5
            )
            output = proc.stdout
            for enc in ["utf-8", "gbk", "gb2312", "latin-1"]:
                try:
                    output = proc.stdout.decode(enc)
                    break
                except (UnicodeDecodeError, LookupError):
                    continue
            success = proc.returncode == 0
            return {"result": output, "success": success, "target": target}
        except Exception as e:
            return {"error": f"Ping failed: {str(e)}"}

    elif ntype == "url_shorten":
        url = _replace_vars(config.get("input_template", ""), context)
        provider = config.get("provider", "tinyurl")
        try:
            import httpx as _httpx
            with _httpx.Client(timeout=10, follow_redirects=True) as client:
                if provider == "tinyurl":
                    resp = client.get(f"https://tinyurl.com/api-create.php", params={"url": url})
                    short_url = resp.text
                elif provider == "is.gd":
                    resp = client.get(f"https://is.gd/create.php", params={"format": "json", "url": url})
                    short_url = resp.json()["shorturl"]
                else:
                    return {"error": f"不支持的缩短服务: {provider}"}
            return {"result": short_url, "original": url}
        except Exception as e:
            return {"error": f"URL 缩短失败: {str(e)}"}

    # ==================== 实用工具类 ====================

    elif ntype == "qrcode_gen":
        input_val = _replace_vars(config.get("input_template", ""), context)
        save_path = config.get("save_path", "")
        size = int(config.get("size", 256))
        try:
            import qrcode
            qr = qrcode.QRCode(version=1, box_size=max(1, size // 256), border=2)
            qr.add_data(input_val)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            if not save_path:
                save_path = os.path.join("generated_media", f"qrcode_{int(__import__('time').time())}.png")
            os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
            img = img.resize((size, size))
            img.save(save_path)
            return {"result": f"二维码已生成", "path": save_path, "content": input_val}
        except ImportError:
            return {"error": "需要安装 qrcode: pip install qrcode[pil]"}
        except Exception as e:
            return {"error": f"二维码生成失败: {str(e)}"}

    elif ntype == "qrcode_scan":
        input_path = _replace_vars(config.get("input_template", ""), context)
        try:
            from PIL import Image
            import pyzbar.pyzbar as pyzbar
            if not input_path or not os.path.exists(input_path):
                return {"error": f"图片不存在: {input_path}"}
            img = Image.open(input_path)
            decoded = pyzbar.decode(img)
            results = [{"data": d.data.decode("utf-8", errors="replace"), "type": d.type} for d in decoded]
            return {"result": results, "count": len(results)}
        except ImportError:
            return {"error": "需要安装 pyzbar: pip install pyzbar"}
        except Exception as e:
            return {"error": f"二维码识别失败: {str(e)}"}

    elif ntype == "pdf_generate":
        input_val = _replace_vars(config.get("input_template", ""), context)
        save_path = config.get("save_path", "")
        title = config.get("title", "")
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            if not save_path:
                save_path = os.path.join("generated_media", f"doc_{int(__import__('time').time())}.pdf")
            os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
            _cn_font = None
            _cn_font_paths = [
                r"C:\Windows\Fonts\simhei.ttf",
                r"C:\Windows\Fonts\msyh.ttc",
                r"C:\Windows\Fonts\simsun.ttc",
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
                "/System/Library/Fonts/PingFang.ttc",
            ]
            for _fp in _cn_font_paths:
                if os.path.exists(_fp):
                    try:
                        pdfmetrics.registerFont(TTFont("CNFont", _fp))
                        _cn_font = "CNFont"
                        break
                    except Exception:
                        continue
            _title_font = _cn_font or "Helvetica-Bold"
            _body_font = _cn_font or "Helvetica"
            c = canvas.Canvas(save_path, pagesize=A4)
            width, height = A4
            y = height - 50
            if title:
                c.setFont(_title_font, 16)
                c.drawString(50, y, title)
                y -= 30
            c.setFont(_body_font, 11)
            for line in input_val.split("\n"):
                if y < 50:
                    c.showPage()
                    y = height - 50
                c.drawString(50, y, line[:100])
                y -= 16
            c.save()
            return {"result": f"PDF 已生成", "path": save_path}
        except ImportError:
            return {"error": "需要安装 reportlab: pip install reportlab"}
        except Exception as e:
            return {"error": f"PDF 生成失败: {str(e)}"}

    elif ntype == "json_diff":
        input_val = _replace_vars(config.get("input_template", ""), context)
        compare_val = _replace_vars(config.get("compare_value", ""), context)
        try:
            if isinstance(input_val, str):
                data1 = json.loads(input_val)
            else:
                data1 = input_val
            if isinstance(compare_val, str):
                data2 = json.loads(compare_val)
            else:
                data2 = compare_val

            def _diff(a, b, path=""):
                differences = []
                if type(a) != type(b):
                    differences.append({"path": path or "/", "type": "type_mismatch", "a": str(type(a).__name__), "b": str(type(b).__name__)})
                    return differences
                if isinstance(a, dict):
                    for key in set(list(a.keys()) + list(b.keys())):
                        if key not in a:
                            differences.append({"path": f"{path}.{key}", "type": "added", "b": b[key]})
                        elif key not in b:
                            differences.append({"path": f"{path}.{key}", "type": "removed", "a": a[key]})
                        else:
                            differences.extend(_diff(a[key], b[key], f"{path}.{key}"))
                elif isinstance(a, list):
                    for i in range(max(len(a), len(b))):
                        if i >= len(a):
                            differences.append({"path": f"{path}[{i}]", "type": "added", "b": b[i]})
                        elif i >= len(b):
                            differences.append({"path": f"{path}[{i}]", "type": "removed", "a": a[i]})
                        else:
                            differences.extend(_diff(a[i], b[i], f"{path}[{i}]"))
                elif a != b:
                    differences.append({"path": path or "/", "type": "changed", "a": a, "b": b})
                return differences

            diffs = _diff(data1, data2)
            return {"result": diffs, "count": len(diffs), "identical": len(diffs) == 0}
        except Exception as e:
            return {"error": f"JSON 对比失败: {str(e)}"}

    elif ntype == "output":
        clean = {}
        for k, v in context.items():
            if k in ("inputs",):
                clean[k] = v
            elif isinstance(v, dict):
                try:
                    json.dumps(v)
                    clean[k] = v
                except (TypeError, ValueError):
                    clean[k] = str(v)
            else:
                clean[k] = v
        return {"output": clean}


    # ==================== 办公自动化类 ====================

    elif ntype == "approval":
        approval_result = {"status": "pending", "approver": "", "comment": ""}
        return {"result": "审批节点已触发", "status": "pending", "note": "需要外部审批系统回调"}

    elif ntype == "email_send":
        smtp_host = config.get("smtp_host", "smtp.qq.com")
        smtp_port = int(config.get("smtp_port", 465))
        smtp_user = config.get("smtp_user", "")
        smtp_pass = config.get("smtp_pass", "")
        use_ssl = config.get("use_ssl", True)
        to_addr = _replace_vars(config.get("to_addr", ""), context)
        subject = _replace_vars(config.get("subject", ""), context)
        body = _replace_vars(config.get("body", "{{input}}"), context)
        is_html = config.get("is_html", False)
        attachments_raw = config.get("attachments", "")
        if not smtp_user or not smtp_pass:
            return {"error": "未配置 SMTP 账号或密码"}
        if not to_addr:
            return {"error": "未配置收件人"}
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.mime.base import MIMEBase
        from email import encoders
        msg = MIMEMultipart()
        msg["From"] = smtp_user
        msg["To"] = to_addr
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "html" if is_html else "plain", "utf-8"))
        attached_files = []
        if attachments_raw:
            paths_str = _replace_vars(attachments_raw, context)
            file_paths = [p.strip() for p in paths_str.replace(";", ",").split(",") if p.strip()]
            for fp in file_paths:
                resolved = fp
                if isinstance(resolved, str) and resolved.startswith(("{", "[")):
                    try:
                        items = json.loads(resolved)
                        if isinstance(items, dict) and "path" in items:
                            resolved = items["path"]
                        elif isinstance(items, list):
                            resolved = items[0].get("path", "") if items and isinstance(items[0], dict) else str(items[0]) if items else ""
                    except (json.JSONDecodeError, TypeError):
                        pass
                if resolved and os.path.exists(resolved):
                    try:
                        part = MIMEBase("application", "octet-stream")
                        with open(resolved, "rb") as f:
                            part.set_payload(f.read())
                        encoders.encode_base64(part)
                        fname = os.path.basename(resolved)
                        part.add_header("Content-Disposition", f'attachment; filename="{fname}"')
                        msg.attach(part)
                        attached_files.append(fname)
                    except Exception as e:
                        return {"error": f"附件添加失败 ({resolved}): {str(e)}"}
                elif resolved:
                    attached_files.append(f"[未找到: {resolved}]")
        try:
            if use_ssl:
                server = smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=15)
            else:
                server = smtplib.SMTP(smtp_host, smtp_port, timeout=15)
                server.starttls()
            server.login(smtp_user, smtp_pass)
            server.sendmail(smtp_user, to_addr.split(","), msg.as_string())
            server.quit()
            return {"result": "邮件已发送", "to": to_addr, "subject": subject, "attachments": attached_files}
        except Exception as e:
            return {"error": f"发送邮件失败: {str(e)}"}

    elif ntype == "excel_write":
        file_path = _replace_vars(config.get("file_path", ""), context)
        data_val = _replace_vars(config.get("data", "{{input}}"), context)
        sheet_name = config.get("sheet_name", "Sheet1")
        try:
            if not file_path:
                file_path = os.path.join("generated_media", f"report_{int(time.time())}.xlsx")
            os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = sheet_name
            rows = json.loads(data_val) if isinstance(data_val, str) else data_val
            if isinstance(rows, list) and len(rows) > 0:
                if isinstance(rows[0], dict):
                    headers = list(rows[0].keys())
                    ws.append(headers)
                    for row in rows:
                        ws.append([row.get(h, "") for h in headers])
                elif isinstance(rows[0], list):
                    for row in rows:
                        ws.append(row)
                else:
                    ws.append([str(r) for r in rows])
            wb.save(file_path)
            wb.close()
            return {"result": "Excel已生成", "path": file_path, "rows": len(rows) if isinstance(rows, list) else 0}
        except Exception as e:
            return {"error": f"Excel写入失败: {str(e)}"}

    elif ntype == "chart_gen":
        data_val = _replace_vars(config.get("data", "{{input}}"), context)
        chart_type = config.get("chart_type", "bar")
        title = config.get("title", "Chart")
        save_path = config.get("save_path", "")
        try:
            import time as _time
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'DejaVu Sans']
            plt.rcParams['axes.unicode_minus'] = False
            rows = json.loads(data_val) if isinstance(data_val, str) else data_val
            if isinstance(rows, list) and len(rows) > 0:
                if isinstance(rows[0], dict):
                    labels = [str(r.get(list(rows[0].keys())[0], i)) for i, r in enumerate(rows)]
                    values = [float(r.get(list(rows[0].keys())[1], 0)) for r in rows]
                elif isinstance(rows[0], list):
                    labels = [str(r[0]) for r in rows]
                    values = [float(r[1]) if len(r) > 1 else 0 for r in rows]
                else:
                    labels = [str(i) for i in range(len(rows))]
                    values = [float(v) for v in rows]
            else:
                return {"error": "数据为空"}
            fig, ax = plt.subplots(figsize=(10, 6))
            if chart_type == "bar":
                ax.bar(labels, values)
            elif chart_type == "line":
                ax.plot(labels, values, marker='o')
            elif chart_type == "pie":
                ax.pie(values, labels=labels, autopct='%1.1f%%')
            elif chart_type == "scatter":
                ax.scatter(range(len(values)), values)
            else:
                ax.bar(labels, values)
            ax.set_title(title)
            plt.tight_layout()
            if not save_path:
                save_path = os.path.join("generated_media", f"chart_{int(_time.time())}.png")
            os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
            fig.savefig(save_path, dpi=150)
            plt.close(fig)
            return {"result": "图表已生成", "path": save_path, "chart_type": chart_type}
        except ImportError:
            return {"error": "需要安装 matplotlib: pip install matplotlib"}
        except Exception as e:
            return {"error": f"图表生成失败: {str(e)}"}

    elif ntype == "statistics":
        data_val = _replace_vars(config.get("data", "{{input}}"), context)
        operation = config.get("operation", "mean")
        try:
            rows = json.loads(data_val) if isinstance(data_val, str) else data_val
            if isinstance(rows, list) and len(rows) > 0:
                if isinstance(rows[0], dict):
                    field = config.get("field", list(rows[0].keys())[0])
                    numbers = [float(r.get(field, 0)) for r in rows if r.get(field) is not None]
                elif isinstance(rows[0], (list, tuple)):
                    numbers = [float(r[0]) for r in rows]
                else:
                    numbers = [float(r) for r in rows]
            else:
                return {"error": "数据为空"}
            import statistics as stats
            ops = {
                "mean": stats.mean, "median": stats.median,
                "stdev": lambda x: stats.stdev(x) if len(x) > 1 else 0,
                "variance": lambda x: stats.variance(x) if len(x) > 1 else 0,
                "min": min, "max": max, "sum": sum, "count": len,
            }
            if operation not in ops:
                return {"error": f"不支持的操作: {operation}"}
            result = ops[operation](numbers)
            return {"result": result, "operation": operation, "count": len(numbers)}
        except Exception as e:
            return {"error": f"统计计算失败: {str(e)}"}

    elif ntype == "calendar_event":
        title = _replace_vars(config.get("title", "会议"), context)
        description = _replace_vars(config.get("description", ""), context)
        start_time = _replace_vars(config.get("start_time", ""), context)
        end_time = _replace_vars(config.get("end_time", ""), context)
        try:
            import time as _time
            import datetime as _dt
            now = _dt.datetime.now()
            if not start_time:
                start_time = now.strftime("%Y-%m-%d %H:%M")
            if not end_time:
                end_time = (now + _dt.timedelta(hours=1)).strftime("%Y-%m-%d %H:%M")
            ics = f"BEGIN:VCALENDAR\nVERSION:2.0\nBEGIN:VEVENT\nSUMMARY:{title}\nDESCRIPTION:{description}\nDTSTART:{start_time.replace('-','').replace(':','').replace(' ','T')}00\nDTEND:{end_time.replace('-','').replace(':','').replace(' ','T')}00\nEND:VEVENT\nEND:VCALENDAR"
            save_path = os.path.join("generated_media", f"event_{int(_time.time())}.ics")
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(ics)
            return {"result": "日历事件已创建", "title": title, "start": start_time, "end": end_time, "ics_path": save_path}
        except Exception as e:
            return {"error": f"创建日历事件失败: {str(e)}"}

    elif ntype == "sentiment_analysis":
        text = _replace_vars(config.get("text", "{{input}}"), context)
        if not llm_router:
            return {"error": "LLM router not available"}
        prompt = f"分析以下文本的情感倾向，只回答一个词：正面/负面/中性\n文本：{text[:2000]}"
        import asyncio, threading
        result_container = [None]
        error_container = [None]
        def _run():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                async def _call():
                    r = ""
                    async for chunk in llm_router.chat_stream([{"role": "user", "content": prompt}], provider=config.get("provider", "agnes")):
                        if chunk == "[DONE]" or chunk.startswith('{"type": "reasoning"'):
                            continue
                        r += chunk
                    return r
                result_container[0] = loop.run_until_complete(_call())
            except Exception as e:
                error_container[0] = str(e)
            finally:
                loop.close()
        t = threading.Thread(target=_run); t.start(); t.join(timeout=60)
        if error_container[0]:
            return {"error": error_container[0]}
        sentiment = result_container[0].strip()
        score = 1.0 if "正面" in sentiment else (-1.0 if "负面" in sentiment else 0.0)
        return {"result": sentiment, "score": score, "text": text[:100]}

    elif ntype == "template_render":
        template_str = config.get("template", "")
        data_val = _replace_vars(config.get("data", "{}"), context)
        try:
            data = json.loads(data_val) if isinstance(data_val, str) else data_val
            result = template_str
            import re
            for key, value in data.items():
                result = re.sub(r"\{\{" + re.escape(key) + "\}\}", str(value), result)
            save_path = config.get("save_path", "")
            if save_path:
                save_path = _replace_vars(save_path, context)
                os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(result)
            return {"result": result, "path": save_path or None}
        except Exception as e:
            return {"error": f"模板渲染失败: {str(e)}"}

    elif ntype == "ssh_exec":
        host = _replace_vars(config.get("host", ""), context)
        username = config.get("username", "")
        password = config.get("password", "")
        command = _replace_vars(config.get("command", "{{input}}"), context)
        port = int(config.get("port", 22))
        if not host or not command:
            return {"error": "未配置主机或命令"}
        try:
            import paramiko
            client = paramiko.SSHClient()
            client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            client.connect(host, port=port, username=username, password=password, timeout=15)
            stdin, stdout, stderr = client.exec_command(command, timeout=30)
            out = stdout.read().decode("utf-8", errors="replace")[:5000]
            err = stderr.read().decode("utf-8", errors="replace")[:2000]
            client.close()
            return {"result": out, "stderr": err, "host": host}
        except ImportError:
            return {"error": "需要安装 paramiko: pip install paramiko"}
        except Exception as e:
            return {"error": f"SSH执行失败: {str(e)}"}

    elif ntype == "log_analyze":
        file_path = _replace_vars(config.get("file_path", "{{input}}"), context)
        keyword = config.get("keyword", "ERROR")
        max_lines = int(config.get("max_lines", 10000))
        try:
            if not file_path or not os.path.exists(file_path):
                return {"error": f"日志文件不存在: {file_path}"}
            matches = []
            total_lines = 0
            keyword_count = 0
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                for i, line in enumerate(f):
                    if i >= max_lines:
                        break
                    total_lines += 1
                    if keyword.lower() in line.lower():
                        keyword_count += 1
                        if len(matches) < 50:
                            matches.append({"line": i + 1, "content": line.strip()[:200]})
            return {"result": {"total_lines": total_lines, "keyword": keyword, "matches": keyword_count, "samples": matches}, "count": keyword_count}
        except Exception as e:
            return {"error": f"日志分析失败: {str(e)}"}

    elif ntype == "backup":
        source = _replace_vars(config.get("source", "{{input}}"), context)
        dest = _replace_vars(config.get("dest", ""), context)
        try:
            if not source or not os.path.exists(source):
                return {"error": f"源路径不存在: {source}"}
            if not dest:
                import datetime as _dt
                dest = os.path.join("generated_media", f"backup_{_dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
            os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
            import shutil
            if os.path.isfile(source):
                shutil.copy2(source, dest)
            else:
                shutil.make_archive(dest.replace(".zip", ""), "zip", source)
                dest = dest if dest.endswith(".zip") else dest + ".zip"
            size = os.path.getsize(dest) if os.path.exists(dest) else 0
            return {"result": "备份完成", "source": source, "dest": dest, "size": size}
        except Exception as e:
            return {"error": f"备份失败: {str(e)}"}

    elif ntype == "data_merge":
        left_key = config.get("left_key", "id")
        right_key = config.get("right_key", "id")
        left_data = context.get(config.get("left_var", "left"), [])
        right_data = context.get(config.get("right_var", "right"), [])
        if isinstance(left_data, str):
            try: left_data = json.loads(left_data)
            except: left_data = []
        if isinstance(right_data, str):
            try: right_data = json.loads(right_data)
            except: right_data = []
        try:
            right_map = {}
            for item in right_data:
                if isinstance(item, dict):
                    right_map[str(item.get(right_key, ""))] = item
            merged = []
            for item in left_data:
                if isinstance(item, dict):
                    r_item = right_map.get(str(item.get(left_key, "")), {})
                    merged.append({**item, **r_item})
                else:
                    merged.append(item)
            return {"result": merged, "count": len(merged)}
        except Exception as e:
            return {"error": f"数据合并失败: {str(e)}"}

    elif ntype == "deduplicate":
        data_val = context.get(config.get("data_var", "input"), [])
        dedup_field = config.get("field", "")
        if isinstance(data_val, str):
            try: data_val = json.loads(data_val)
            except: data_val = [data_val]
        try:
            seen = set()
            result = []
            for item in data_val:
                key = str(item.get(dedup_field, "")) if isinstance(item, dict) and dedup_field else str(item)
                if key not in seen:
                    seen.add(key)
                    result.append(item)
            return {"result": result, "count": len(result), "removed": len(data_val) - len(result)}
        except Exception as e:
            return {"error": f"去重失败: {str(e)}"}

    elif ntype == "pivot_table":
        data_val = context.get(config.get("data_var", "input"), [])
        group_by = config.get("group_by", "")
        agg_field = config.get("agg_field", "")
        agg_func = config.get("agg_func", "sum")
        if isinstance(data_val, str):
            try: data_val = json.loads(data_val)
            except: data_val = []
        try:
            groups = {}
            for item in data_val:
                if isinstance(item, dict):
                    key = str(item.get(group_by, "default"))
                    val = float(item.get(agg_field, 0))
                    groups.setdefault(key, []).append(val)
            result = []
            for key, vals in groups.items():
                agg_val = sum(vals) if agg_func == "sum" else (len(vals) if agg_func == "count" else (max(vals) if agg_func == "max" else (min(vals) if agg_func == "min" else sum(vals)/len(vals))))
                result.append({group_by: key, f"{agg_func}({agg_field})": agg_val, "count": len(vals)})
            return {"result": result, "groups": len(result)}
        except Exception as e:
            return {"error": f"透视表失败: {str(e)}"}

    elif ntype == "correlation":
        x_data = context.get(config.get("x_var", "x"), [])
        y_data = context.get(config.get("y_var", "y"), [])
        if isinstance(x_data, str):
            try: x_data = json.loads(x_data)
            except: x_data = []
        if isinstance(y_data, str):
            try: y_data = json.loads(y_data)
            except: y_data = []
        try:
            x = [float(v) for v in x_data]
            y = [float(v) for v in y_data]
            if len(x) != len(y) or len(x) < 2:
                return {"error": "数据长度不一致或不足2个"}
            import statistics as stats
            mean_x, mean_y = stats.mean(x), stats.mean(y)
            cov = sum((xi - mean_x) * (yi - mean_y) for xi, yi in zip(x, y)) / (len(x) - 1)
            std_x, std_y = stats.stdev(x), stats.stdev(y)
            r = cov / (std_x * std_y) if std_x * std_y != 0 else 0
            return {"result": r, "correlation": r, "strength": "强" if abs(r) > 0.7 else ("中等" if abs(r) > 0.4 else "弱"), "count": len(x)}
        except Exception as e:
            return {"error": f"相关性分析失败: {str(e)}"}

    elif ntype == "docx_read":
        file_path = _replace_vars(config.get("file_path", "{{input}}"), context)
        try:
            if not file_path or not os.path.exists(file_path):
                return {"error": f"文件不存在: {file_path}"}
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return {"result": "\n".join(paragraphs), "paragraphs": len(paragraphs)}
        except ImportError:
            return {"error": "需要安装 python-docx: pip install python-docx"}
        except Exception as e:
            return {"error": f"读取Word失败: {str(e)}"}

    elif ntype == "docx_write":
        save_path = _replace_vars(config.get("save_path", ""), context)
        title = _replace_vars(config.get("title", ""), context)
        content = _replace_vars(config.get("content", "{{input}}"), context)
        try:
            from docx import Document
            if not save_path:
                save_path = os.path.join("generated_media", f"doc_{int(time.time())}.docx")
            os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
            doc = Document()
            if title:
                doc.add_heading(title, level=1)
            for para in content.split("\n"):
                if para.strip():
                    doc.add_paragraph(para)
            doc.save(save_path)
            return {"result": "Word文档已生成", "path": save_path}
        except ImportError:
            return {"error": "需要安装 python-docx: pip install python-docx"}
        except Exception as e:
            return {"error": f"生成Word失败: {str(e)}"}

    elif ntype == "encrypt":
        text = _replace_vars(config.get("text", "{{input}}"), context)
        action = config.get("action", "encrypt")
        key = config.get("key", "")
        try:
            import time as _time
            import hashlib, base64
            if action == "encrypt":
                if not key:
                    key = base64.b64encode(os.urandom(32)).decode()
                h = hashlib.sha256((text + key).encode()).hexdigest()
                encoded = base64.b64encode(text.encode()).decode()
                return {"result": encoded, "hash": h, "key": key, "action": "encrypt"}
            else:
                result = base64.b64decode(text.encode()).decode()
                return {"result": result, "action": "decrypt"}
        except Exception as e:
            return {"error": f"加密失败: {str(e)}"}

    elif ntype == "jwt_generate":
        payload_data = _replace_vars(config.get("payload", "{}"), context)
        secret = config.get("secret", "default-secret")
        expire_minutes = int(config.get("expire_minutes", 30))
        try:
            import hashlib, base64, json as _json
            payload = json.loads(payload_data) if isinstance(payload_data, str) else payload_data
            import datetime as _dt
            payload["exp"] = (_dt.datetime.utcnow() + _dt.timedelta(minutes=expire_minutes)).isoformat()
            payload["iat"] = _dt.datetime.utcnow().isoformat()
            header = base64.b64encode(_json.dumps({"alg": "HS256", "typ": "JWT"}).encode()).decode().rstrip("=")
            body = base64.b64encode(_json.dumps(payload).encode()).decode().rstrip("=")
            sig = base64.b64encode(hashlib.sha256(f"{header}.{body}.{secret}".encode()).digest()).decode().rstrip("=")
            token = f"{header}.{body}.{sig}"
            return {"result": token, "expires_in": expire_minutes}
        except Exception as e:
            return {"error": f"JWT生成失败: {str(e)}"}

    elif ntype == "schedule_trigger":
        schedule_type = config.get("schedule_type", "interval")
        interval = int(config.get("interval", 60))
        return {"result": "调度已注册", "type": schedule_type, "interval": interval, "note": "调度触发器已注册，需要外部调度器执行"}

    return {"result": "unknown node type"}


_versions_col = None


def _get_versions_col():
    global _versions_col
    if _versions_col is None:
        from pymongo import MongoClient
        MONGO_URL = os.environ.get("MONGODB_URL", "mongodb://ai_mongo:ai_mongo_2024@localhost:27017")
        client = MongoClient(MONGO_URL, serverSelectionTimeoutMS=3000)
        db = client["ai_framework"]
        _versions_col = db["workflow_versions"]
    return _versions_col


def export_workflow(workflow_id: str) -> Optional[Dict]:
    wf = get_workflow(workflow_id)
    if not wf:
        return None
    return {
        "name": wf.get("name", ""),
        "description": wf.get("description", ""),
        "nodes": wf.get("nodes", []),
        "edges": wf.get("edges", []),
        "exported_at": datetime.now().isoformat(),
        "version": "1.0"
    }


def import_workflow(data: dict) -> Dict:
    import_data = {
        "name": data.get("name", "导入的工作流"),
        "description": data.get("description", ""),
        "nodes": data.get("nodes", []),
        "edges": data.get("edges", []),
        "status": "draft"
    }
    return create_workflow(import_data)


def save_version(workflow_id: str, comment: str = "") -> Optional[Dict]:
    wf = get_workflow(workflow_id)
    if not wf:
        return None
    col = _get_versions_col()
    versions_count = col.count_documents({"workflow_id": workflow_id})
    version_doc = {
        "workflow_id": workflow_id,
        "version": versions_count + 1,
        "name": wf.get("name", ""),
        "description": wf.get("description", ""),
        "nodes": wf.get("nodes", []),
        "edges": wf.get("edges", []),
        "status": wf.get("status", "draft"),
        "comment": comment,
        "created_at": datetime.now().isoformat()
    }
    col.insert_one(version_doc)
    return {"version": version_doc["version"], "created_at": version_doc["created_at"]}


def list_versions(workflow_id: str) -> List[Dict]:
    col = _get_versions_col()
    results = []
    for doc in col.find({"workflow_id": workflow_id}).sort("version", -1):
        doc.pop("_id", None)
        results.append(doc)
    return results


def get_version(workflow_id: str, version: int) -> Optional[Dict]:
    col = _get_versions_col()
    doc = col.find_one({"workflow_id": workflow_id, "version": version})
    if doc:
        doc.pop("_id", None)
    return doc


def restore_version(workflow_id: str, version: int) -> bool:
    v = get_version(workflow_id, version)
    if not v:
        return False
    update_data = {
        "name": v.get("name", ""),
        "description": v.get("description", ""),
        "nodes": v.get("nodes", []),
        "edges": v.get("edges", []),
    }
    return update_workflow(workflow_id, update_data)


def delete_version(workflow_id: str, version: int) -> bool:
    col = _get_versions_col()
    result = col.delete_one({"workflow_id": workflow_id, "version": version})
    return result.deleted_count > 0


def run_workflow(workflow_id: str, inputs: Dict[str, Any] = None, llm_router=None) -> WorkflowRun:
    wf_data = get_workflow(workflow_id)
    if not wf_data:
        raise ValueError(f"Workflow {workflow_id} not found")

    run = WorkflowRun(workflow_id=workflow_id, inputs=inputs or {})
    save_run(run)

    nodes = wf_data.get("nodes", [])
    edges = wf_data.get("edges", [])

    if not nodes:
        run.status = "completed"
        run.finished_at = datetime.now().isoformat()
        update_run(run)
        return run

    levels = _topological_sort(nodes, edges)
    node_map = {n["id"]: n for n in nodes}
    context = {"inputs": inputs or {}}
    if isinstance(inputs, dict):
        for k, v in inputs.items():
            if k not in context:
                context[k] = v

    for level in levels:
        for node_id in level:
            node = node_map.get(node_id)
            if not node:
                continue
            try:
                result = _execute_node(node, context, llm_router)
                context[node_id] = result
                if isinstance(result, dict):
                    for k, v in result.items():
                        if k not in context:
                            context[k] = v
                run.node_results[node_id] = {"status": "completed", "result": result}
            except Exception as e:
                run.node_results[node_id] = {"status": "failed", "error": str(e)}
                run.status = "failed"
                run.error = str(e)
                run.finished_at = datetime.now().isoformat()
                update_run(run)
                return run

    clean_outputs = {}
    for k, v in context.items():
        if isinstance(v, dict):
            try:
                json.dumps(v)
                clean_outputs[k] = v
            except (TypeError, ValueError):
                clean_outputs[k] = str(v)
        elif isinstance(v, list):
            try:
                json.dumps(v)
                clean_outputs[k] = v
            except (TypeError, ValueError):
                clean_outputs[k] = str(v)
        else:
            clean_outputs[k] = v
    run.outputs = clean_outputs
    run.status = "completed"
    run.finished_at = datetime.now().isoformat()
    update_run(run)
    return run
